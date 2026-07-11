"""Extracts info from pdf/docx/doc resumes using OpenAI and returns JSON.

Usage: python resume_parser.py resume.pdf
"""

import json
import os
import re
import subprocess
import sys
import tempfile
from dotenv import load_dotenv
load_dotenv()
import fitz
from docx import Document
from openai import OpenAI

MODEL = os.getenv("PARSER_MODEL", "gpt-4o")

PROMPT = """Read this resume and think step by step before answering:
1. Find the contact details (name, email, phone).
2. Find every education entry (institution, degree, graduation year).
3. Find every job (company, position, description, duration). Resumes list jobs
   in different date formats and column layouts, so check the whole text carefully.
4. Find the skills list.
5. Find certifications and projects, if any.

Return ONLY valid JSON in this format:
{
  "reasoning": "your short step-by-step notes from above",
  "contact_information": {"name": "", "email": "", "phone": ""},
  "education": [{"institution": "", "degree": "", "graduation_year": ""}],
  "work_experience": [{"company": "", "position": "", "description": "", "duration": ""}],
  "skills": [],
  "certifications": [],
  "projects": [{"name": "", "description": ""}]
}

Rules:
- Only use values that are actually in the resume, never guess.
- Use null for missing fields, [] for empty lists.
- Keep descriptions to 1-2 lines each.

Resume:
"""


def extract_text(path):
    ext = os.path.splitext(path)[1].lower()

    if ext == ".pdf":
        doc = fitz.open(path)
        text = "\n".join(page.get_text() for page in doc)
        doc.close()

    elif ext == ".docx":
        doc = Document(path)
        lines = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                lines.append(" | ".join(c.text.strip() for c in row.cells))
        text = "\n".join(lines)

    elif ext == ".doc":
        outdir = tempfile.mkdtemp()
        subprocess.run(["soffice", "--headless", "--convert-to", "docx",
                        "--outdir", outdir, path], check=True, capture_output=True)
        name = os.path.basename(path).rsplit(".", 1)[0] + ".docx"
        return extract_text(os.path.join(outdir, name))

    else:
        raise ValueError(f"unsupported file type: {ext}")

    if not text.strip():
        raise ValueError("no text could be extracted from the file")
    return text


def flag_unmatched(label, values, plain_text, warnings):
    for value in values:
        normalized = " ".join((value or "").lower().split())
        if normalized and normalized not in plain_text:
            warnings.append(f"{label} not found in resume: {value}")


def clean_result(data, text):
    """Drop or flag values the model may have made up."""
    warnings = []
    contact = data.get("contact_information") or {}

    email = contact.get("email")
    if email and (not re.match(r"^[\w.+-]+@[\w-]+\.[\w.-]+$", email) or email not in text):
        contact["email"] = None
        warnings.append(f"dropped email not found in resume: {email}")

    phone = contact.get("phone")
    if phone:
        digits = re.sub(r"\D", "", phone)
        if len(digits) < 7 or digits not in re.sub(r"\D", "", text):
            contact["phone"] = None
            warnings.append(f"dropped phone not found in resume: {phone}")

    plain = " ".join(text.lower().split())
    companies = [job.get("company") for job in data.get("work_experience") or []]
    projects = [p.get("name") for p in data.get("projects") or []]
    flag_unmatched("company", companies, plain, warnings)
    flag_unmatched("certification", data.get("certifications") or [], plain, warnings)
    flag_unmatched("project", projects, plain, warnings)

    data["warnings"] = warnings
    return data


def parse_resume(path):
    text = extract_text(path)
    client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
    response = client.chat.completions.create(
        model=MODEL,
        temperature=0,
        response_format={"type": "json_object"},
        messages=[{"role": "user", "content": PROMPT + text[:15000]}],
    )
    data = json.loads(response.choices[0].message.content)
    data.pop("reasoning", None)
    return clean_result(data, text)


if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("usage: python resume_parser.py <resume file>")
        sys.exit(1)
    print(json.dumps(parse_resume(sys.argv[1]), indent=2, ensure_ascii=False))
